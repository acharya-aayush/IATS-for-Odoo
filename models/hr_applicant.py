import base64
import re
from io import BytesIO
from zipfile import ZipFile

from odoo import api, fields, models
from odoo.exceptions import UserError

try:
    import pdfplumber
except Exception:  # pragma: no cover - optional dependency
    pdfplumber = None

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency
    Document = None

try:
    from striprtf.striprtf import rtf_to_text
except Exception:  # pragma: no cover - optional dependency
    rtf_to_text = None

try:
    from PIL import Image
except Exception:  # pragma: no cover - optional dependency
    Image = None

try:
    import pytesseract
except Exception:  # pragma: no cover - optional dependency
    pytesseract = None


class HrApplicant(models.Model):
    _inherit = "hr.applicant"

    iats_profile_id = fields.Many2one("iats.job.profile", string="IATS Profile", index=True, copy=False)
    iats_state = fields.Selection( 
        [
            ("pending", "Pending"),
            ("ready", "Ready"),
            ("scored", "Scored"),
            ("failed", "Failed"),
        ],
        default="pending",
        copy=False,
        tracking=True,
    )
    iats_recommendation = fields.Selection( 
        [
            ("shortlist", "Shortlist"),
            ("review", "Review"),
            ("reject", "Reject"),
        ],
        copy=False,
        tracking=True,
    )
    iats_score = fields.Float(string="Overall Score", copy=False, tracking=True)
    iats_keyword_score = fields.Float(string="Keyword Score", copy=False)
    iats_skill_score = fields.Float(string="Skill Score", copy=False)
    iats_experience_score = fields.Float(string="Experience Score", copy=False)
    iats_education_score = fields.Float(string="Education Score", copy=False)
    iats_completeness_score = fields.Float(string="Completeness Score", copy=False)
    iats_last_screened_at = fields.Datetime(string="Last Screened", copy=False)
    iats_years_experience = fields.Float(string="Years of Experience", copy=False)
    iats_matched_keywords = fields.Char(string="Matched Keywords", copy=False)
    iats_match_summary = fields.Text(string="Match Summary", copy=False)

    iats_resume_file = fields.Binary(string="IATS Resume", attachment=True)
    iats_resume_filename = fields.Char(string="IATS Resume Filename", copy=False)
    iats_resume_text = fields.Text(string="Extracted Text", compute="_compute_iats_resume_content", store=True)
    iats_resume_parse_status = fields.Selection( 
        [
            ("missing", "Missing"),
            ("ready", "Ready"),
            ("empty", "No Readable Text"),
            ("failed", "Extraction Failed"),
        ],
        compute="_compute_iats_resume_content",
        store=True,
        copy=False,
    )
    iats_resume_parse_message = fields.Char(string="Parse Message", compute="_compute_iats_resume_content", store=True, copy=False)

    def _get_default_iats_resume_filename(self):
        self.ensure_one()
        base_name = re.sub(r"[^A-Za-z0-9._-]+", "_", (self.partner_name or "resume").strip()).strip("_")
        return f"{base_name or 'resume'}.bin"

    def _ensure_iats_profile(self):
        for applicant in self:
            if not applicant.iats_profile_id and applicant.job_id:
                profile = applicant.job_id.sudo()._get_or_create_iats_profile()
                applicant.with_context(skip_iats_refresh=True).write({
                    "iats_profile_id": profile.id,
                })

    def _sync_iats_resume_from_latest_attachment(self):
        attachment_model = self.env["ir.attachment"].sudo()
        supported_ext = (".pdf", ".docx", ".doc", ".rtf", ".txt", ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".odt")
        for applicant in self:
            if applicant.iats_resume_file:
                continue
            attachments = attachment_model.search(
                [("res_model", "=", "hr.applicant"), ("res_id", "=", applicant.id), ("res_field", "=", False)],
                order="create_date desc, id desc",
            )
            attachment = attachments.filtered(
                lambda record: (
                    (record.name and record.name.lower().endswith(supported_ext))
                    or (record.mimetype and ("image/" in record.mimetype or "pdf" in record.mimetype or "word" in record.mimetype or "text/" in record.mimetype))
                    or (not record.name or "." not in record.name)
                )
            )[:1] or attachments[:1]
            if attachment and attachment.datas:
                applicant.with_context(skip_iats_refresh=True).write({
                    "iats_resume_file": attachment.datas,
                    "iats_resume_filename": attachment.name or applicant.iats_resume_filename or applicant._get_default_iats_resume_filename(),
                })

    def _ensure_iats_resume_filename(self):
        for applicant in self:
            if applicant.iats_resume_file and not applicant.iats_resume_filename:
                applicant.with_context(skip_iats_refresh=True).write({
                    "iats_resume_filename": applicant._get_default_iats_resume_filename(),
                })

    def _reset_iats_scores(self):
        self.with_context(skip_iats_refresh=True).write({
            "iats_state": "pending",
            "iats_recommendation": False,
            "iats_score": 0.0,
            "iats_keyword_score": 0.0,
            "iats_skill_score": 0.0,
            "iats_experience_score": 0.0,
            "iats_education_score": 0.0,
            "iats_completeness_score": 0.0,
            "iats_last_screened_at": False,
            "iats_years_experience": 0.0,
            "iats_matched_keywords": False,
            "iats_match_summary": False,
        })

    def _extract_text_from_pdf(self, resume_bytes):
        if not pdfplumber:
            return ""
        with pdfplumber.open(BytesIO(resume_bytes)) as pdf:
            return "\n".join((page.extract_text() or "") for page in pdf.pages)

    def _extract_text_from_docx(self, resume_bytes):
        if not Document:
            return ""
        document = Document(BytesIO(resume_bytes))
        chunks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = (cell.text or "").strip()
                    if text:
                        chunks.append(text)
        return "\n".join(chunks)

    def _extract_text_from_plain(self, resume_bytes):
        for encoding in ("utf-8", "utf-16", "utf-16le", "utf-16be", "cp1252", "latin-1"):
            try:
                return resume_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        return resume_bytes.decode("latin-1", errors="ignore")

    def _extract_text_from_rtf(self, resume_bytes):
        if not rtf_to_text:
            return ""
        return rtf_to_text(resume_bytes.decode("latin-1", errors="ignore"))

    def _extract_text_from_odt(self, resume_bytes):
        try:
            with ZipFile(BytesIO(resume_bytes)) as archive:
                content = archive.read("content.xml")
        except Exception:
            return ""
        return re.sub(r"<[^>]+>", " ", content.decode("utf-8", errors="ignore"))

    def _extract_text_with_ocr(self, resume_bytes):
        if not Image or not pytesseract:
            return ""
        try:
            image = Image.open(BytesIO(resume_bytes))
            return pytesseract.image_to_string(image)
        except Exception:
            return ""

    def _extract_iats_resume_text(self, resume_bytes):
        if resume_bytes.startswith(b"%PDF-"):
            return self._extract_text_from_pdf(resume_bytes)
        if resume_bytes.startswith(b"PK\x03\x04"):
            text = self._extract_text_from_docx(resume_bytes)
            return text or self._extract_text_from_odt(resume_bytes)
        if resume_bytes.lstrip().startswith((b"{\\rtf", b"{\\RTF")):
            return self._extract_text_from_rtf(resume_bytes)
        if resume_bytes.startswith((b"\x89PNG", b"\xff\xd8\xff", b"GIF8", b"II*\x00", b"MM\x00*")):
            return self._extract_text_with_ocr(resume_bytes)
        if b"\x00" not in resume_bytes:
            return self._extract_text_from_plain(resume_bytes)
        return ""

    @api.depends("iats_resume_file")
    def _compute_iats_resume_content(self):
        for applicant in self:
            if not applicant.iats_resume_file:
                applicant.iats_resume_text = False
                applicant.iats_resume_parse_status = "missing"
                applicant.iats_resume_parse_message = False
                continue
            try:
                resume_bytes = base64.b64decode(applicant.iats_resume_file)
                text = applicant._extract_iats_resume_text(resume_bytes)

                if text:
                    text = re.sub(r"(?<!\n)([•●▪])", r"\n\1 ", text)
                    headings = ["WORK EXPERIENCE", "EDUCATION", "SKILLS", "PROJECTS", "CERTIFICATIONS", "VOLUNTEERING", "LEADERSHIP", "SUMMARY"]
                    for heading in headings:
                        text = re.sub(rf"((?<!\n)\b{heading}\b)", rf"\n\n\1\n", text, flags=re.IGNORECASE)

                cleaned_text = re.sub(r"[ \t]+", " ", text or "").strip()
                cleaned_text = re.sub(r"\n{3,}", "\n\n", cleaned_text)
                
                if cleaned_text:
                    applicant.iats_resume_text = cleaned_text
                    applicant.iats_resume_parse_status = "ready"
                    applicant.iats_resume_parse_message = False
                else:
                    applicant.iats_resume_text = False
                    applicant.iats_resume_parse_status = "empty"
                    applicant.iats_resume_parse_message = "No readable text was extracted from the uploaded resume."
            except Exception:
                applicant.iats_resume_text = False
                applicant.iats_resume_parse_status = "failed"
                applicant.iats_resume_parse_message = "Resume parsing failed for the uploaded file."

    @api.model_create_multi
    def create(self, vals_list):
        applicants = super().create(vals_list)
        for applicant in applicants:
            applicant._after_iats_resume_change()
        return applicants

    def write(self, vals):
        result = super().write(vals)
        if self.env.context.get("skip_iats_refresh"):
            return result
        if any(field in vals for field in ("job_id", "iats_profile_id", "iats_resume_file", "attachment_ids")):
            self._reset_iats_scores()
            for applicant in self:
                applicant._after_iats_resume_change()
        return result

    def _after_iats_resume_change(self):
        for applicant in self:
            applicant._ensure_iats_profile()
            applicant._sync_iats_resume_from_latest_attachment()
            applicant._ensure_iats_resume_filename()
            applicant._compute_iats_resume_content()

            if applicant.iats_resume_text and not applicant.linkedin_profile:
                linked_match = re.search(r"linkedin\.com/in/[\w\-]+", applicant.iats_resume_text, flags=re.IGNORECASE)
                if linked_match:
                    applicant.with_context(skip_iats_refresh=True).write({
                        "linkedin_profile": f"https://www.{linked_match.group(0).lower()}"
                    })

            if applicant.iats_resume_parse_status == "ready":
                applicant.with_context(skip_iats_refresh=True).write({"iats_state": "ready"})
                if applicant.iats_profile_id and applicant.iats_profile_id.auto_screen_enabled:
                    try:
                        self.env.ref("iats_recruitment.ir_cron_iats_profile_screening")._trigger()
                    except Exception:
                        pass
            elif applicant.iats_resume_parse_status in ("empty", "failed"):
                applicant.with_context(skip_iats_refresh=True).write({"iats_state": "failed"})
            else:
                applicant.with_context(skip_iats_refresh=True).write({"iats_state": "pending"})

    def _apply_iats_stage_routing(self):
        self.ensure_one()
        profile = self.iats_profile_id
        if not profile or not profile.auto_move_stage:
            return
        if self.iats_recommendation == "shortlist" and profile.shortlist_stage_id:
            self.stage_id = profile.shortlist_stage_id.id
        elif self.iats_recommendation == "review" and profile.review_stage_id:
            self.stage_id = profile.review_stage_id.id
        elif self.iats_recommendation == "reject" and profile.reject_stage_id:
            self.stage_id = profile.reject_stage_id.id

    def _run_iats_screening(self, profile=False, force=False):
        for applicant in self:
            current_profile = profile or applicant.iats_profile_id
            if not current_profile:
                applicant._ensure_iats_profile()
                current_profile = applicant.iats_profile_id
            if not current_profile or applicant.iats_resume_parse_status != "ready" or not applicant.iats_resume_text:
                if force:
                    raise UserError("Applicant resume is not ready for IATS screening.")
                continue
            scores = current_profile._score_applicant(applicant)
            applicant.with_context(skip_iats_refresh=True).write({
                "iats_profile_id": current_profile.id,
                "iats_state": "scored",
                "iats_recommendation": scores["recommendation"],
                "iats_score": scores["total_score"],
                "iats_keyword_score": scores["keyword_score"],
                "iats_skill_score": scores["skill_score"],
                "iats_experience_score": scores["experience_score"],
                "iats_education_score": scores["education_score"],
                "iats_completeness_score": scores["completeness_score"],
                "iats_last_screened_at": fields.Datetime.now(),
                "iats_years_experience": scores["years_experience"],
                "iats_matched_keywords": scores["matched_keywords"],
                "iats_match_summary": scores["summary"],
            })
            applicant._apply_iats_stage_routing()
            if current_profile.auto_notify_reviewers and scores["recommendation"] in ("shortlist", "review"):
                current_profile._notify_reviewers(applicant)
        return True

    def action_run_iats_screening(self):
        self._run_iats_screening(force=True)
        return {
            "type": "ir.actions.client",
            "tag": "display_notification",
            "params": {
                "title": "IATS Screening Complete",
                "message": "Applicant screening has been refreshed.",
                "type": "success",
                "sticky": False,
            },
        }
